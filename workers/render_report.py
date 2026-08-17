from __future__ import annotations

import argparse
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt

from report_common import parse_markdown
from report_pdf import render_pdf


def archive_existing(output_dir: Path) -> None:
    archive_dir = output_dir / "archive"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    current = output_dir / "meeting.docx"
    if not current.exists():
        return
    archive_dir.mkdir(parents=True, exist_ok=True)
    target = archive_dir / f"meeting_{timestamp}.docx"
    counter = 2
    while target.exists():
        target = archive_dir / f"meeting_{timestamp}_{counter}.docx"
        counter += 1
    shutil.move(str(current), str(target))


def add_docx_paragraph(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith(("- ", "* ")):
        document.add_paragraph(stripped[2:].strip(), style="List Bullet")
    elif stripped[0:2].isdigit() and stripped[2:4] in (". ", ") "):
        document.add_paragraph(stripped[4:].strip(), style="List Number")
    else:
        document.add_paragraph(stripped)


def render_docx(markdown: str, output_path: Path) -> None:
    title, blocks = parse_markdown(markdown)
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = document.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "맑은 고딕"
    styles["Title"].font.size = Pt(22)
    for name in ("Heading 1", "Heading 2"):
        styles[name].font.name = "맑은 고딕"

    document.add_heading(title, level=0)
    for block in blocks:
        document.add_heading(block.title, level=min(max(block.level - 1, 1), 2))
        for line in block.lines:
            add_docx_paragraph(document, line)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="meeting.md를 DOCX/PDF로 렌더링합니다.")
    parser.add_argument("--input", required=True, help="입력 meeting.md 경로")
    parser.add_argument("--output-dir", required=True, help="출력 디렉터리")
    parser.add_argument("--pdf-name", default="meeting.pdf", help="최종 PDF 파일명")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    output_dir = Path(args.output_dir).resolve()
    if not input_path.exists():
        raise SystemExit(f"입력 파일을 찾을 수 없습니다: {input_path}")
    pdf_name = args.pdf_name
    if Path(pdf_name).name != pdf_name or not pdf_name.lower().endswith(".pdf"):
        raise SystemExit(f"올바르지 않은 PDF 파일명입니다: {pdf_name}")
    output_dir.mkdir(parents=True, exist_ok=True)
    markdown = input_path.read_text(encoding="utf-8")

    archive_existing(output_dir)
    for existing_pdf in output_dir.glob("*.pdf"):
        existing_pdf.unlink()
    render_docx(markdown, output_dir / "meeting.docx")
    pdf_path = output_dir / pdf_name
    render_pdf(markdown, pdf_path, Path(__file__).resolve().parents[2])
    print(f"DOCX 생성: {output_dir / 'meeting.docx'}")
    print(f"PDF 생성: {pdf_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
